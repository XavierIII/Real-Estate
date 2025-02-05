import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import requests
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import GridSearchCV
from transformers import pipeline
from fbprophet import Prophet
import spacy
from docx import Document
import streamlit as st
import pyttsx3
import speech_recognition as sr
import time
from sklearn.metrics.pairwise import cosine_similarity

# Load NLP model
nlp = spacy.load("en_core_web_sm")

# Real-time Data Integration (Zillow or other API)
def fetch_live_property_data(city, state, min_price, max_price):
    # Example API endpoint - Replace with your actual API and API key
    url = f'https://api.zillow.com/v2/properties'
    params = {
        'city': city,
        'state': state,
        'min_price': min_price,
        'max_price': max_price,
        'api_key': 'YOUR_ZILLOW_API_KEY'  # Replace with your actual API key
    }
    response = requests.get(url, params=params)
    if response.status_code == 200:
        properties = response.json()
        return properties['property_list']
    else:
        return None

# Define AI-powered property search and comparison
def train_investment_model():
    # Historical data: Features [Price, Square Footage, Location Score], Target: ROI
    historical_data = np.array([[200000, 1500, 8], [250000, 1800, 9], [300000, 2000, 10]])
    roi = np.array([12, 15, 18])  # Example ROI values
    model = RandomForestRegressor()
    model.fit(historical_data, roi)
    return model

def recommend_properties(user_input, properties_database):
    similarities = cosine_similarity([user_input], properties_database)
    recommended_idx = similarities.argsort()[0][-5:]  # Top 5 similar properties
    return recommended_idx

def compare_properties(properties):
    comparison_data = {}
    for property in properties:
        price_per_sqft = property['price'] / property['sqft']
        comparison_data[property['address']] = {
            'Price per SqFt': price_per_sqft,
            'ROI': property['roi'],
            'DSCR': property['dscr']
        }
    return comparison_data

def plot_price_trends(price_data):
    dates = [item['date'] for item in price_data]
    prices = [item['price'] for item in price_data]
    plt.plot(dates, prices)
    plt.title('Price Trend Over Time')
    plt.xlabel('Date')
    plt.ylabel('Price')
    st.pyplot(plt)

def calculate_coc(initial_investment, annual_rent, loan_payment):
    return (annual_rent - loan_payment) / initial_investment * 100

def sensitivity_analysis(initial_investment, annual_rent, vacancy_rate, appreciation_rate):
    rent_adjusted = annual_rent * (1 - vacancy_rate)
    price_appreciation = initial_investment * (1 + appreciation_rate)
    roi = (rent_adjusted / price_appreciation) * 100
    return roi

def calculate_risk(location_score, property_age, market_volatility):
    risk_score = (location_score * 0.4) + (property_age * 0.3) + (market_volatility * 0.3)
    return risk_score

def estimate_repair_cost(sqft, property_age):
    base_cost = 5  # $5 per square foot
    age_factor = 1.2 if property_age > 20 else 1.0
    return sqft * base_cost * age_factor

def get_neighborhood_data(zip_code):
    neighborhood_info = {
        "crime_rate": 0.2,
        "school_rating": 8,
        "walkability": 0.8
    }
    return neighborhood_info

def financing_options(loan_amount, interest_rate, term_years):
    monthly_interest_rate = interest_rate / 12 / 100
    number_of_payments = term_years * 12
    monthly_payment = (loan_amount * monthly_interest_rate) / (1 - (1 + monthly_interest_rate) ** -number_of_payments)
    return monthly_payment

def send_property_alert(properties):
    for prop in properties:
        print(f"Alert: New property found - {prop['address']}")

def calculate_tax(property_value, tax_rate=0.01):
    return property_value * tax_rate

def depreciation_schedule(property_value, years=27.5):
    annual_depreciation = property_value / years
    return annual_depreciation

def wholesaling_profit(purchase_price, repair_cost, sale_price):
    profit = sale_price - purchase_price - repair_cost
    return profit

def equity_financing(investment, equity_share):
    return investment * equity_share

def long_term_projection(property_value, years, appreciation_rate):
    future_value = property_value * (1 + appreciation_rate) ** years
    return future_value

def fetch_property_images(property_id):
    images = ["image1.jpg", "image2.jpg"]
    return images

def crm_integration(client_data):
    print(f"Sending data to CRM: {client_data}")

def fetch_market_news():
    url = "https://newsapi.org/v2/everything?q=real%20estate&apiKey=YOUR_NEWSAPI_KEY"
    response = requests.get(url)
    return response.json()

def save_user_preferences(user_preferences):
    print(f"Saving preferences: {user_preferences}")

def generate_legal_documents(property_data):
    doc = Document()
    doc.add_paragraph(f"Agreement for Property: {property_data['address']}")
    doc.add_paragraph(f"Price: ${property_data['price']}")
    doc.save(f"{property_data['address']}_contract.docx")

def share_on_social_media(property_data):
    print(f"Sharing property {property_data['address']} on social media.")

def export_report(properties):
    df = pd.DataFrame(properties)
    df.to_excel("property_report.xlsx", index=False)

def assign_user_roles(user, role):
    print(f"Assigning {role} role to user {user}")

# Voice Command Integration
recognizer = sr.Recognizer()

def listen_for_command():
    with sr.Microphone() as source:
        print("Listening for command...")
        audio = recognizer.listen(source)

    try:
        command = recognizer.recognize_google(audio)
        print(f"Recognized Command: {command}")
        return command
    except sr.UnknownValueError:
        print("Sorry, I did not understand that.")
        return None
    except sr.RequestError:
        print("Sorry, there was a problem with the speech recognition service.")
        return None

# Model Performance Tuning using GridSearchCV
def tune_model():
    # Larger dataset for better prediction
    historical_data = np.array([[200000, 1500, 8], [250000, 1800, 9], [300000, 2000, 10], [350000, 2200, 7], [400000, 2400, 10]])
    roi = np.array([12, 15, 18, 10, 20])

    model = RandomForestRegressor()
    param_grid = {
        'n_estimators': [100, 200, 300],
        'max_depth': [None, 10, 20, 30],
        'min_samples_split': [2, 5, 10],
    }

    grid_search = GridSearchCV(estimator=model, param_grid=param_grid, cv=5, n_jobs=-1)
    grid_search.fit(historical_data, roi)

    best_model = grid_search.best_estimator_
    print("Best Model:", best_model)
    return best_model

# Streamlit UI Setup
st.title("AI-Powered Real Estate Investment Tool")

# User Inputs for Property Search Criteria
state = st.text_input("Enter State")
city = st.text_input("Enter City")
min_price = st.number_input("Min Price", min_value=0)
max_price = st.number_input("Max Price", min_value=0)

# Button to trigger the property search
if st.button("Search Properties"):
    # Fetch live data from API
    properties = fetch_live_property_data(city, state, min_price, max_price)
    
    if properties:
        comparison = compare_properties(properties)
        st.write("Property Comparison:", comparison)

        # Plot historical price trends
        price_data = [{'date': '2023-01-01', 'price': 250000}, {'date': '2023-02-01', 'price': 255000}]
        plot_price_trends(price_data)
        
        # Sensitivity analysis example
        roi_sensitivity = sensitivity_analysis(200000, 24000, 0.05, 0.03)
        st.write(f"ROI after sensitivity adjustments: {roi_sensitivity}%")

        # Financing options
        monthly_payment = financing_options(200000, 5, 30)
        st.write(f"Monthly Payment: ${monthly_payment}")

        # Calculate repair costs
        repair_cost = estimate_repair_cost(1500, 25)
        st.write(f"Estimated Repair Cost: ${repair_cost}")

        # Tax & Depreciation
        tax = calculate_tax(250000)
        depreciation = depreciation_schedule(250000)
        st.write(f"Annual Tax: ${tax}, Annual Depreciation: ${depreciation}")

        # Share property on social media
        share_on_social_media(properties[0])

        # Export the report
        export_report(properties)
    else:
        st.write("No properties found.")

# Run model tuning
best_model = tune_model()

# Example: Legal Document Generation
example_property = {'address': '123 Main St', 'price': 350000}
generate_legal_documents(example_property)

# Listen for a voice command
command = listen_for_command()

